"""Generate a standalone DOCX with the cleaned reference list, ready to paste into VKR."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

OUT = Path(r"C:\Users\Admin\Desktop\Diplom\docs\vkr\Список_источников_FINAL.docx")

SOURCES: list[str] = [
    "1. Стратегия развития туризма Иркутской области на период до 2030 года [Электронный ресурс] : утверждена распоряжением Правительства Иркутской области. — URL: https://irkobl.ru/sites/tour/ (дата обращения: 15.03.2026).",
    "2. Исполатов В. П. Анализ программных интерфейсов российских агрегаторов бронирования средств размещения : материалы исследования в рамках проекта / В. П. Исполатов. – Иркутск, 2025.",
    "3. Иванов В. В. Гостиничный менеджмент : учебное пособие / В. В. Иванов, А. Б. Волов. – Москва : ИНФРА-М, 2007. – 384 с.",
    "4. Паспорт туристского кластера Иркутской области [Электронный ресурс] : утверждён Агентством по туризму Иркутской области. — URL: https://irkobl.ru/sites/tour/utp.php (дата обращения: 25.03.2026).",
    "5. STR Global. Аналитика индустрии гостеприимства [Электронный ресурс]. — URL: https://str.com/ (дата обращения: 18.03.2026).",
    "6. Yandex DataLens. Документация сервиса визуализации данных [Электронный ресурс]. — URL: https://cloud.yandex.ru/docs/datalens/ (дата обращения: 20.03.2026).",
    "7. Xotelo. Hotel Price Monitor and API [Электронный ресурс]. — URL: https://xotelo.com/ (дата обращения: 22.03.2026).",
    "8. Open-Meteo. Бесплатный API погодных данных [Электронный ресурс]. — URL: https://open-meteo.com/ (дата обращения: 22.03.2026).",
    "9. FastAPI. Документация фреймворка [Электронный ресурс]. — URL: https://fastapi.tiangolo.com/ (дата обращения: 20.03.2026).",
    "10. PostgreSQL 16 Documentation [Электронный ресурс]. — URL: https://www.postgresql.org/docs/16/ (дата обращения: 20.03.2026).",
    "11. asyncpg. Асинхронный драйвер PostgreSQL для Python [Электронный ресурс]. — URL: https://pypi.org/project/asyncpg/ (дата обращения: 20.03.2026).",
    "12. Redis. Документация хранилища данных в памяти [Электронный ресурс]. — URL: https://redis.io/docs/ (дата обращения: 22.03.2026).",
    "13. nginx. Документация веб-сервера [Электронный ресурс]. — URL: https://nginx.org/ru/docs/ (дата обращения: 21.03.2026).",
    "14. APScheduler. Документация библиотеки планирования задач [Электронный ресурс]. — URL: https://apscheduler.readthedocs.io/ (дата обращения: 19.03.2026).",
    "15. OWASP. API Security Top 10 [Электронный ресурс]. — URL: https://owasp.org/API-Security/ (дата обращения: 23.03.2026).",
    "16. State of JavaScript 2025. Ежегодное исследование экосистемы JavaScript [Электронный ресурс]. — URL: https://stateofjs.com/ (дата обращения: 20.03.2026).",
    "17. Tufte E. R. The Visual Display of Quantitative Information / E. R. Tufte. – 2nd ed. – Cheshire : Graphics Press, 2001. – 197 p.",
    "18. Tufte E. R. Envisioning Information / E. R. Tufte. – Cheshire : Graphics Press, 1990. – 126 p.",
    "19. Few S. Information Dashboard Design: Displaying Data for At-a-Glance Monitoring / S. Few. – 2nd ed. – Burlingame : Analytics Press, 2013. – 260 p.",
    "20. Norman D. A. The Design of Everyday Things / D. A. Norman. – Revised and Expanded ed. – New York : Basic Books, 2013. – 368 p.",
    "21. Nielsen J. 10 Usability Heuristics for User Interface Design [Электронный ресурс] // Nielsen Norman Group. — URL: https://www.nngroup.com/articles/ten-usability-heuristics/ (дата обращения: 25.03.2026).",
    "22. Frost B. Atomic Design [Электронный ресурс] : методология проектирования интерфейсных систем // Atomic Design. — URL: https://atomicdesign.bradfrost.com/ (дата обращения: 25.03.2026).",
    "23. Shneiderman B. The Eyes Have It: A Task by Data Type Taxonomy for Information Visualizations / B. Shneiderman // Proceedings of the IEEE Symposium on Visual Languages. – Boulder, 1996. – P. 336–343.",
    "24. Vite. Документация сборщика [Электронный ресурс]. — URL: https://vite.dev/ (дата обращения: 20.03.2026).",
    "25. Tailwind CSS. Документация CSS-фреймворка [Электронный ресурс]. — URL: https://tailwindcss.com/docs (дата обращения: 20.03.2026).",
    "26. Recharts. Библиотека графиков для React [Электронный ресурс]. — URL: https://recharts.org/ (дата обращения: 21.03.2026).",
    "27. Belobaba P. P. Air Travel Demand and Airline Seat Inventory Management [Электронный ресурс] : PhD diss. / P. P. Belobaba ; Massachusetts Institute of Technology. – Cambridge, 1987. — URL: https://dspace.mit.edu/handle/1721.1/14800 (дата обращения: 06.05.2026).",
    "28. Kimes S. E. Revenue Management: A Retrospective / S. E. Kimes // Cornell Hotel and Restaurant Administration Quarterly. – 2003. – Vol. 44, No. 5. – P. 131–138.",
    "29. Cross R. G. Revenue Management: Hard-Core Tactics for Market Domination / R. G. Cross. – New York : Broadway Books, 1997. – 288 p.",
    "30. Time Series Analysis: Forecasting and Control / G. E. P. Box [и др.]. – 5th ed. – Hoboken : Wiley, 2015. – 712 p.",
    "31. Taylor S. J. Forecasting at Scale / S. J. Taylor, B. Letham // The American Statistician. – 2018. – Vol. 72, No. 1. – P. 37–45.",
    "32. NeuralProphet: Explainable Forecasting at Scale [Электронный ресурс] / O. Triebe [и др.]. – 2021. – Препринт arXiv:2111.15397. — URL: https://arxiv.org/abs/2111.15397 (дата обращения: 25.03.2026).",
    "33. Chen T. XGBoost: A Scalable Tree Boosting System / T. Chen, C. Guestrin // Proceedings of the 22nd ACM SIGKDD International Conference on Knowledge Discovery and Data Mining. – San Francisco, 2016. – P. 785–794.",
    "34. Dietterich T. G. Ensemble Methods in Machine Learning / T. G. Dietterich // Proceedings of the First International Workshop on Multiple Classifier Systems (MCS 2000). – Cagliari, 2000. – P. 1–15.",
    "35. LangGraph. Документация фреймворка для построения агентов на графах состояний [Электронный ресурс]. — URL: https://www.langchain.com/langgraph (дата обращения: 22.03.2026).",
    "36. Groq. Платформа быстрого инференса LLM [Электронный ресурс]. — URL: https://groq.com/ (дата обращения: 22.03.2026).",
    "37. Mistral AI. Документация API [Электронный ресурс]. — URL: https://docs.mistral.ai/ (дата обращения: 23.03.2026).",
    "38. Competing Against Luck: The Story of Innovation and Customer Choice / C. M. Christensen [и др.]. – New York : HarperBusiness, 2016. – 288 p.",
    "39. Wedel M. Market Segmentation: Conceptual and Methodological Foundations / M. Wedel, W. A. Kamakura. – 2nd ed. – Norwell : Kluwer Academic Publishers, 2000. – 382 p.",
    "40. Федеральная служба государственной статистики (Росстат). Статистика туризма по Иркутской области [Электронный ресурс]. — URL: https://rosstat.gov.ru/ (дата обращения: 15.03.2026).",
    "41. STR Global Hotel Benchmarking Glossary. Методические материалы CoStar Group [Электронный ресурс]. — URL: https://str.com/data-insights/glossary (дата обращения: 06.05.2026).",
    "42. AirDNA. Аналитическая платформа данных краткосрочной аренды [Электронный ресурс]. — URL: https://www.airdna.co/ (дата обращения: 25.03.2026).",
    "43. Электронная путёвка. Единая информационная система электронных путёвок (ЕИСЭП) [Электронный ресурс]. — URL: https://ev.economy.gov.ru/ (дата обращения: 25.03.2026).",
    "44. React. Документация библиотеки [Электронный ресурс]. — URL: https://react.dev/ (дата обращения: 20.03.2026).",
    "45. Docker Documentation [Электронный ресурс]. — URL: https://docs.docker.com/ (дата обращения: 21.03.2026).",
    "46. SQLAlchemy 2.0. Документация ORM [Электронный ресурс]. — URL: https://docs.sqlalchemy.org/en/20/ (дата обращения: 19.03.2026).",
    "47. aiohttp. Документация библиотеки асинхронных HTTP-запросов для Python [Электронный ресурс]. — URL: https://docs.aiohttp.org/ (дата обращения: 18.03.2026).",
    "48. 101Hotels.com. Агрегатор средств размещения [Электронный ресурс]. — URL: https://www.101hotels.com/ (дата обращения: 25.03.2026).",
    "49. Zheng A. Feature Engineering for Machine Learning: Principles and Techniques for Data Scientists / A. Zheng, A. Casari. – Sebastopol : O'Reilly Media, 2018. – 215 p.",
    "50. ChromaDB. Документация векторной базы данных [Электронный ресурс]. — URL: https://docs.trychroma.com/ (дата обращения: 21.03.2026).",
    "51. Rosenfeld L. Information Architecture for the World Wide Web / L. Rosenfeld, P. Morville. – 3rd ed. – Sebastopol : O'Reilly Media, 2006. – 504 p.",
    "52. TanStack Query. Документация библиотеки управления серверным состоянием для React [Электронный ресурс]. — URL: https://tanstack.com/query/ (дата обращения: 22.03.2026).",
    "53. Apache ECharts. Библиотека визуализации данных [Электронный ресурс]. — URL: https://echarts.apache.org/ (дата обращения: 21.03.2026).",
]


def _set_run_font(run, name: str = "Times New Roman", size_pt: int = 14, bold: bool = False) -> None:
    run.font.name = name
    run.font.size = Pt(size_pt)
    run.bold = bold
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        from docx.oxml import OxmlElement
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rFonts.set(qn(attr), name)


def main() -> None:
    doc = Document()

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(14)

    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(1.5)

    heading = doc.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.paragraph_format.space_before = Pt(0)
    heading.paragraph_format.space_after = Pt(12)
    heading.paragraph_format.line_spacing = 1.5
    h_run = heading.add_run("СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ")
    _set_run_font(h_run, bold=True)

    for entry in SOURCES:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.first_line_indent = Cm(1.25)
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(entry)
        _set_run_font(run, bold=False)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Saved: {OUT}")
    print(f"Entries: {len(SOURCES)}")


if __name__ == "__main__":
    main()
