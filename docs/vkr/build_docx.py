"""
Build OTCHET_PO_PRAKTIKE.docx from .md with exact styles from coursework.

Styles extracted from "Курсовая работа Исполатов.docx":
  - "Обычный текст": TNR 14pt, justify, indent 1.27cm, line spacing 1.5
  - "Заголовок первого уровня": inherits "Обычный текст", bold, center, indent 0
  - "Заголовок второго уровня": inherits "Обычный текст", bold, justify, indent 1.27cm
  - "Подпись рисунка": inherits "Обычный текст", center, indent 0
  - "List Paragraph": inherits Normal, left_indent 1.27cm
  - Margins: left 3cm, right 1cm, top 2cm, bottom 2cm
"""

import re
import sys
import io
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Cm, Pt, RGBColor, Emu
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")

SCRIPT_DIR = Path(__file__).parent
MD_SOURCE = SCRIPT_DIR / "OTCHET_PO_PRAKTIKE.md"
DOCX_OUTPUT = SCRIPT_DIR / "OTCHET_PO_PRAKTIKE.docx"

FONT = "Times New Roman"


def _set_rfonts(rpr, name=FONT):
    """Set all font families on an rPr element."""
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rfonts.set(qn(attr), name)


def _create_styles(doc):
    """Create exact copies of coursework styles."""
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(12)
    _set_rfonts(normal.element.get_or_add_rPr())

    body = doc.styles.add_style("Обычный текст", WD_STYLE_TYPE.PARAGRAPH)
    body.quick_style = True
    body.priority = 1
    body.font.name = FONT
    body.font.size = Pt(14)
    body.font.bold = None
    body.font.italic = None
    body.font.color.rgb = RGBColor(0, 0, 0)
    _set_rfonts(body.element.get_or_add_rPr())
    pf = body.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.first_line_indent = Cm(1.27)
    pf.line_spacing = 1.5
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)

    h1 = doc.styles.add_style("Заголовок первого уровня", WD_STYLE_TYPE.PARAGRAPH)
    h1.quick_style = True
    h1.priority = 2
    h1.base_style = body
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0, 0, 0)
    _set_rfonts(h1.element.get_or_add_rPr())
    h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h1.paragraph_format.first_line_indent = Cm(0)
    h1.paragraph_format.space_before = Pt(0)
    h1.paragraph_format.space_after = Pt(0)
    ppr1 = h1.element.get_or_add_pPr()
    lvl1 = OxmlElement("w:outlineLvl")
    lvl1.set(qn("w:val"), "0")
    ppr1.append(lvl1)

    h2 = doc.styles.add_style("Заголовок второго уровня", WD_STYLE_TYPE.PARAGRAPH)
    h2.quick_style = True
    h2.priority = 3
    h2.base_style = body
    h2.font.bold = True
    _set_rfonts(h2.element.get_or_add_rPr())
    h2.paragraph_format.space_before = Pt(0)
    h2.paragraph_format.space_after = Pt(0)
    ppr2 = h2.element.get_or_add_pPr()
    lvl2 = OxmlElement("w:outlineLvl")
    lvl2.set(qn("w:val"), "1")
    ppr2.append(lvl2)

    caption = doc.styles.add_style("Подпись рисунка", WD_STYLE_TYPE.PARAGRAPH)
    caption.quick_style = True
    caption.priority = 4
    caption.base_style = body
    caption.font.color.rgb = RGBColor(0, 0, 0)
    _set_rfonts(caption.element.get_or_add_rPr())
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.first_line_indent = Cm(0)

    return {"body": body, "h1": h1, "h2": h2, "caption": caption}


def _add_run(paragraph, text, bold=False, italic=False):
    """Add a run with proper font settings."""
    run = paragraph.add_run(text)
    run.font.name = FONT
    run.font.size = Pt(14)
    if bold:
        run.font.bold = True
    if italic:
        run.font.italic = True
    _set_rfonts(run.element.get_or_add_rPr())
    return run


def _process_inline(paragraph, text, base_bold=False):
    """Parse markdown inline formatting (**bold**, *italic*) and add runs."""
    pattern = re.compile(r'(\*\*\*(.+?)\*\*\*|\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`)')
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            _add_run(paragraph, text[pos:m.start()], bold=base_bold)
        if m.group(2):
            _add_run(paragraph, m.group(2), bold=True, italic=True)
        elif m.group(3):
            _add_run(paragraph, m.group(3), bold=True)
        elif m.group(4):
            _add_run(paragraph, m.group(4), italic=True)
        elif m.group(5):
            _add_run(paragraph, m.group(5))
        pos = m.end()
    if pos < len(text):
        _add_run(paragraph, text[pos:], bold=base_bold)


def _add_table(doc, header_row, data_rows, styles):
    """Add a table with proper styling."""
    ncols = len(header_row)
    table = doc.add_table(rows=1 + len(data_rows), cols=ncols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    for i, cell_text in enumerate(header_row):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_run(p, cell_text.strip(), bold=True)

    for r, row_data in enumerate(data_rows):
        for c, cell_text in enumerate(row_data):
            cell = table.rows[r + 1].cells[c]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            _process_inline(p, cell_text.strip())

    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.0
                p.paragraph_format.first_line_indent = Cm(0)
                for run in p.runs:
                    run.font.size = Pt(12)
                    run.font.name = FONT
                    _set_rfonts(run.element.get_or_add_rPr())


def _parse_table_block(lines, start):
    """Parse a markdown table starting at `start`. Returns (header, rows, end_idx)."""
    header_line = lines[start]
    cols = [c.strip() for c in header_line.strip().strip("|").split("|")]

    sep_idx = start + 1
    if sep_idx >= len(lines) or not re.match(r'^[\s|:-]+$', lines[sep_idx]):
        return None, None, start

    rows = []
    idx = sep_idx + 1
    while idx < len(lines) and "|" in lines[idx] and not lines[idx].strip().startswith("#"):
        row = [c.strip() for c in lines[idx].strip().strip("|").split("|")]
        while len(row) < len(cols):
            row.append("")
        rows.append(row[:len(cols)])
        idx += 1

    return cols, rows, idx


def build():
    """Main build function."""
    md_text = MD_SOURCE.read_text(encoding="utf-8")
    lines = md_text.split("\n")

    doc = Document()

    section = doc.sections[0]
    section.left_margin = Cm(3)
    section.right_margin = Cm(1)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)

    styles = _create_styles(doc)
    doc.add_paragraph().clear()

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        if stripped.startswith("| ") and i + 1 < len(lines) and re.match(r'^[\s|:-]+$', lines[i + 1]):
            header, rows, end_idx = _parse_table_block(lines, i)
            if header and rows:
                _add_table(doc, header, rows, styles)
                i = end_idx
                continue
            i += 1
            continue

        if stripped.startswith("# "):
            text = stripped[2:].strip()
            p = doc.add_paragraph(style=styles["h1"])
            _add_run(p, text, bold=True)
            i += 1
            continue

        if stripped.startswith("## "):
            text = stripped[3:].strip()
            p = doc.add_paragraph(style=styles["h2"])
            _add_run(p, text, bold=True)
            i += 1
            continue

        if stripped.startswith("### ") or stripped.startswith("#### "):
            level = 3 if stripped.startswith("### ") else 4
            text = stripped[level + 1:].strip()
            p = doc.add_paragraph(style=styles["h2"])
            _add_run(p, text, bold=True)
            i += 1
            continue

        img_match = re.match(r'^!\[.*?\]\((.+?)\)', stripped)
        if img_match:
            img_rel = img_match.group(1)
            img_path = SCRIPT_DIR / img_rel
            if img_path.exists():
                p = doc.add_paragraph(style=styles["caption"])
                p.paragraph_format.space_before = Pt(6)
                run = p.add_run()
                run.add_picture(str(img_path), width=Cm(15))
            i += 1
            continue

        if re.match(r'^\*?\[.*рисун', stripped, re.IGNORECASE) or \
           re.match(r'^Рисунок\s+\d', stripped):
            p = doc.add_paragraph(style=styles["caption"])
            _process_inline(p, stripped.strip("*[]"))
            i += 1
            continue

        caption_match = re.match(r'^(Таблица\s+\d+\s*[–—-]\s*.+)', stripped)
        if caption_match:
            p = doc.add_paragraph(style=styles["body"])
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.first_line_indent = Cm(0)
            _process_inline(p, caption_match.group(1))
            i += 1
            continue

        list_match = re.match(r'^(\d+)\.\s+(.+)', stripped)
        if list_match:
            num = list_match.group(1)
            text = list_match.group(2)
            p = doc.add_paragraph(style=styles["body"])
            _process_inline(p, f"{num}. {text}")
            i += 1
            continue

        if stripped.startswith("- "):
            text = stripped[2:]
            p = doc.add_paragraph(style=styles["body"])
            _process_inline(p, f"– {text}")
            i += 1
            continue

        p = doc.add_paragraph(style=styles["body"])
        _process_inline(p, stripped)
        i += 1

    if doc.paragraphs and not doc.paragraphs[0].text:
        p0 = doc.paragraphs[0]._element
        p0.getparent().remove(p0)

    doc.save(str(DOCX_OUTPUT))
    size = DOCX_OUTPUT.stat().st_size
    print(f"OK: {DOCX_OUTPUT.name} ({size:,} bytes)")


if __name__ == "__main__":
    build()
