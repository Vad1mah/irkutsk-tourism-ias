"""Verify styles in generated docx match coursework."""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")

from pathlib import Path
from docx import Document

def emu_to_cm(emu):
    return round(emu / 914400 * 2.54, 2) if emu else None

def emu_to_pt(emu):
    return round(emu / 12700, 1) if emu else None

doc = Document(str(Path(__file__).parent / "OTCHET_PO_PRAKTIKE.docx"))

sec = doc.sections[0]
print(f"Margins: L={emu_to_cm(sec.left_margin)} R={emu_to_cm(sec.right_margin)} "
      f"T={emu_to_cm(sec.top_margin)} B={emu_to_cm(sec.bottom_margin)}")

used = set()
for p in doc.paragraphs:
    used.add(p.style.name)
print(f"\nUsed styles: {sorted(used)}")

for sname in sorted(used):
    s = doc.styles[sname]
    pf = s.paragraph_format
    f = s.font
    base = s.base_style.name if s.base_style else "-"
    print(f"\n  [{sname}] base={base}")
    print(f"    font={f.name} size={emu_to_pt(f.size)}pt bold={f.bold}")
    print(f"    align={pf.alignment} indent={emu_to_cm(pf.first_line_indent)}cm "
          f"spacing={pf.line_spacing}")

print(f"\nFirst 20 paragraphs:")
for i, p in enumerate(doc.paragraphs[:20]):
    txt = p.text[:80]
    rinfo = ""
    if p.runs:
        r = p.runs[0]
        rinfo = f" font={r.font.name} {emu_to_pt(r.font.size)}pt b={r.font.bold}"
    print(f"  [{i}] style={p.style.name:30s}{rinfo} | {txt}")
