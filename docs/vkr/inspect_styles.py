"""Inspect all styles used in the coursework docx."""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, Emu

def emu_to_cm(emu):
    if emu is None:
        return None
    return round(emu / 914400 * 2.54, 2)

def emu_to_pt(emu):
    if emu is None:
        return None
    return round(emu / 12700, 1)

import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")

src = list(Path(__file__).parent.joinpath("cursovaya_backup").glob("*.docx"))[0]
doc = Document(str(src))

print("=" * 80)
print(f"FILE: {src.name}")
print("=" * 80)

sec = doc.sections[0]
print(f"\nSECTION margins:")
print(f"  left:   {emu_to_cm(sec.left_margin)} cm")
print(f"  right:  {emu_to_cm(sec.right_margin)} cm")
print(f"  top:    {emu_to_cm(sec.top_margin)} cm")
print(f"  bottom: {emu_to_cm(sec.bottom_margin)} cm")
print(f"  page width:  {emu_to_cm(sec.page_width)} cm")
print(f"  page height: {emu_to_cm(sec.page_height)} cm")

print(f"\n{'=' * 80}")
print("DEFINED STYLES (used):")
print("=" * 80)

used_styles = set()
for p in doc.paragraphs:
    used_styles.add(p.style.name)
    for r in p.runs:
        if r.style and r.style.name != "Default Paragraph Font":
            used_styles.add(f"  [run] {r.style.name}")

for t in doc.tables:
    for row in t.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                used_styles.add(f"  [table] {p.style.name}")

for sname in sorted(used_styles):
    print(f"  - {sname}")

print(f"\n{'=' * 80}")
print("STYLE DETAILS (paragraph styles):")
print("=" * 80)

seen = set()
for p in doc.paragraphs:
    sname = p.style.name
    if sname in seen:
        continue
    seen.add(sname)
    s = p.style
    pf = s.paragraph_format
    f = s.font
    
    print(f"\n--- {sname} ---")
    print(f"  base_style: {s.base_style.name if s.base_style else None}")
    print(f"  font.name: {f.name}")
    print(f"  font.size: {emu_to_pt(f.size) if f.size else None} pt")
    print(f"  font.bold: {f.bold}")
    print(f"  font.italic: {f.italic}")
    print(f"  font.color.rgb: {f.color.rgb if f.color and f.color.rgb else None}")
    print(f"  alignment: {pf.alignment}")
    print(f"  first_line_indent: {emu_to_cm(pf.first_line_indent)} cm")
    print(f"  left_indent: {emu_to_cm(pf.left_indent)} cm")
    print(f"  line_spacing: {pf.line_spacing}")
    print(f"  space_before: {emu_to_pt(pf.space_before) if pf.space_before else None} pt")
    print(f"  space_after: {emu_to_pt(pf.space_after) if pf.space_after else None} pt")

print(f"\n{'=' * 80}")
print("PARAGRAPH-BY-PARAGRAPH (first 80):")
print("=" * 80)

for i, p in enumerate(doc.paragraphs[:80]):
    text = p.text[:90].replace("\n", " ")
    run_info = ""
    if p.runs:
        r = p.runs[0]
        rsize = emu_to_pt(r.font.size) if r.font.size else "-"
        rbold = "B" if r.font.bold else ""
        rital = "I" if r.font.italic else ""
        rfont = r.font.name or "-"
        run_info = f" | run: {rfont} {rsize}pt {rbold}{rital}"
    
    pf = p.paragraph_format
    align = str(pf.alignment).split(".")[-1] if pf.alignment else "inherit"
    indent = emu_to_cm(pf.first_line_indent) if pf.first_line_indent else "-"
    spacing = pf.line_spacing if pf.line_spacing else "-"
    
    print(f"  [{i:3d}] style={p.style.name:25s} align={align:10s} indent={str(indent):6s} spacing={str(spacing):5s}{run_info} | {text}")
